"""Build abstrak Portal HC KPB .docx mirroring paten DCS format.

Mirrors:
- A4 page with line numbering (every 5 lines) in left margin
- Centered page number, "Abstrak" label, bold uppercase letter-spaced title
- Justified body, Times New Roman 12pt, line-spacing 2.0
- First-line indent
- Italic emphasis on English terms
"""
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

OUT = "docs/abstract/abstrak-portal-hc-kpb.docx"

doc = Document()

# Page setup A4 with margins (extra left for line numbers)
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Mm(25)
section.bottom_margin = Mm(25)
section.left_margin = Mm(35)
section.right_margin = Mm(22)

# Enable line numbering: every 5 lines, continuous, distance from text
sectPr = section._sectPr
ln = OxmlElement('w:lnNumType')
ln.set(qn('w:countBy'), '5')
ln.set(qn('w:start'), '1')
ln.set(qn('w:restart'), 'continuous')
ln.set(qn('w:distance'), '360')  # 360 twips ~= 6.35mm
sectPr.append(ln)


def set_font(run, name="Times New Roman", size=12, bold=False, italic=False,
             letter_spacing=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    # East Asia font set so name applies everywhere
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), name)
    if letter_spacing is not None:
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:val'), str(int(letter_spacing * 20)))  # twentieths of pt
        rPr.append(spacing)


def add_centered(text, size=12, bold=False, letter_spacing=None,
                 space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    # Suppress line numbering on header lines
    pPr = p._element.get_or_add_pPr()
    suppress = OxmlElement('w:suppressLineNumbers')
    pPr.append(suppress)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, letter_spacing=letter_spacing)
    return p


# Page number "1"
add_centered("1", size=12, space_after=18)
# "Abstrak" label
add_centered("Abstrak", size=12, space_after=8)
# Title
add_centered(
    "SISTEM DIGITAL PROGRAM PENGEMBANGAN KOMPETENSI OPERASI KILANG "
    "PROFESIONAL (PROTON)",
    size=12,
    bold=True,
    letter_spacing=2,
    space_after=18,
)

# Body paragraph: justified, first-line indent, line spacing 2.0
body = doc.add_paragraph()
body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
body.paragraph_format.first_line_indent = Mm(12)
body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
body.paragraph_format.space_after = Pt(0)
body.paragraph_format.space_before = Pt(0)


def add_run(text, italic=False):
    run = body.add_run(text)
    set_font(run, size=12, italic=italic)


# Body content — segments split for italic English terms
# Each tuple: (text, italic_flag)
segments = [
    ("Invensi ini berkaitan dengan suatu sistem digital berbasis web "
     "(ASP.NET Core MVC) yang mengoperasionalkan program PROTON (", False),
    ("Professional Refinery Operations Competency Development", True),
    ("), yaitu program pengembangan kompetensi terstruktur berprinsip "
     "SMART bagi pekerja kilang minyak pada fase operasi. Program "
     "dijalankan pada dua jalur keahlian (Panelman dan Operator) selama "
     "tiga tahun berjenjang — ", False),
    ("Foundation", True),
    (", Pendalaman, dan ", False),
    ("Mastery", True),
    (" — melalui dua modul terintegrasi: ", False),
    ("Competency Management Platform", True),
    (" (CMP) menyusun Kebutuhan Kompetensi Jabatan (KKJ), melaksanakan "
     "asesmen kompetensi teknis, dan menyepakati Silabus, sedangkan ", False),
    ("Competency Development Platform", True),
    (" (CDP) mengeksekusinya dengan metode ", False),
    ("blended learning", True),
    (". Pelaksanaan PROTON mengikuti alur enam langkah: (1) penetapan "
     "Silabus per jabatan berbasis KKJ dan penugasan ", False),
    ("coach", True),
    (" oleh ", False),
    ("Human Capital", True),
    ("; (2) pengunggahan bukti ", False),
    ("deliverable", True),
    (" oleh ", False),
    ("coachee", True),
    (" via fitur ", False),
    ("Individual Development Plan", True),
    (" (IDP); (3) pencatatan sesi pendampingan via fitur ", False),
    ("Coaching", True),
    (" PROTON berpedoman ", False),
    ("Coaching Guidance", True),
    (" per dimensi kompetensi; (4) persetujuan berlapis oleh ", False),
    ("Sr Supervisor", True),
    (", ", False),
    ("Section Head", True),
    (", dan ", False),
    ("review", True),
    (" final ", False),
    ("Human Capital", True),
    ("; (5) ", False),
    ("Final Assessment", True),
    (" berupa ujian pilihan ganda otomatis (Tahun 1–2) atau wawancara "
     "panel (Tahun 3), ambang kelulusan ditetapkan (75); serta (6) "
     "perekaman permanen pada Histori PROTON. Setelah lulus, sertifikat "
     "digital terbit otomatis dengan masa berlaku dan rantai pembaruan "
     "lintas-periode, sehingga status kompetensi dan riwayat "
     "pengembangan termutakhirkan secara konsisten. Dengan demikian, "
     "invensi mereplikasi program PROTON dalam lingkungan digital yang "
     "terstruktur, terukur, dan terlacak.", False),
]

for text, italic in segments:
    add_run(text, italic=italic)

doc.save(OUT)
print(f"Saved: {OUT}")
