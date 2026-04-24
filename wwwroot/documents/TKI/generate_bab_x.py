"""
Generator BAB X INSTRUKSI KERJA + BAB XI INDIKATOR KEBERHASILAN
untuk Draft TKI Penggunaan HC PRIME — PT Kilang Pertamina Balikpapan

Output:
  - Draft-BAB-X-INSTRUKSI-KERJA.docx  (template TKI-GAST-003: header tabel, footer, numbering Roman)
  - Draft-BAB-X-INSTRUKSI-KERJA.html  (preview untuk browser)
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent
SHOT = HERE / "screenshots"

# =========================================================
# STRUKTUR DATA — BAB X
# Setiap item: (heading_level, text, role_callout_or_None, screenshot_or_None)
# Role callout: list of (role_tag, step_text, screenshot_filename or None)
# =========================================================

OUTLINE = {
    "A. AKSES": {
        "items": [
            ("1. Login ke HC Prime", [
                ("USER", "Buka URL http://10.55.3.3/KPB-PortalHC melalui browser (Chrome/Edge/Firefox).", "A01-login-page.png"),
                ("USER", "Masukkan email Pertamina (nama.lengkap@pertamina.com) dan password.", None),
                ("USER", "Klik tombol \"Login\".", None),
                ("USER", "Bila kredensial valid, sistem mengarahkan ke halaman Beranda sesuai peran pengguna.", "A02-beranda.png"),
            ]),
        ]
    },
    "B. MODUL CMP — COMPETENCY MANAGEMENT": {
        "items": [
            ("1. Melihat KKJ Pribadi (Kebutuhan Kompetensi Jabatan)", [
                ("USER", "Pada halaman Beranda, klik menu CMP.", "B01-cmp-dashboard.png"),
                ("USER", "Klik kartu \"Dokumen KKJ\". Sistem menampilkan daftar KKJ sesuai jabatan pengguna.", None),
                ("USER", "Klik ikon unduh untuk mendapatkan dokumen matriks KKJ (PDF/Excel).", "B02-dokumen-kkj.png"),
            ]),
            ("2. Melihat Matriks KKJ Bagian (Atasan)", [
                ("ATASAN", "Buka menu CMP > Dokumen KKJ.", None),
                ("ATASAN", "Pilih filter bagian (NGP / GAST / RFCC / DHT & HMU / Utilities II). Sistem menampilkan matriks kompetensi seluruh pekerja di bagian tersebut.", None),
                ("ATASAN", "Klik salah satu baris untuk melihat detail gap kompetensi per pekerja.", None),
            ]),
            ("3. Mengikuti Assessment Online", [
                ("USER", "Setelah mendapat notifikasi penugasan assessment, klik menu CMP > Assessment.", "B03-assessment-list.png"),
                ("USER", "Sistem menampilkan daftar paket assessment yang ditugaskan dengan status Open.", None),
                ("USER", "Klik tombol \"Mulai\" pada paket yang hendak dikerjakan.", None),
                ("USER", "Baca instruksi dan ketentuan, lalu klik \"Setuju & Mulai\".", None),
                ("USER", "Kerjakan soal satu per satu. Timer berjalan di pojok kanan atas.", None),
                ("USER", "Tipe soal yang tersedia: (1) Pilihan ganda; (2) Multiple answer; (3) Essay; (4) Mix.", None),
                ("USER", "Gunakan tombol \"Sebelumnya\" / \"Selanjutnya\" untuk navigasi antar soal.", None),
                ("USER", "Bila seluruh soal telah dijawab, klik \"Submit\". Sistem meminta konfirmasi.", None),
                ("USER", "Sistem menyimpan jawaban, melakukan auto-grading untuk tipe non-essay, lalu menampilkan halaman Ringkasan Ujian.", None),
                ("USER", "Untuk soal Essay: skor pending sampai HC melakukan grading manual.", None),
            ]),
            ("4. Melihat Rekam Jejak Assessment Pribadi (Records)", [
                ("USER", "Klik menu CMP > Records.", "B04-records.png"),
                ("USER", "Sistem menampilkan tabel seluruh riwayat assessment pribadi (kategori, nama paket, tanggal, skor, status).", None),
                ("USER", "Klik ikon mata pada salah satu baris untuk melihat detail hasil per soal.", None),
                ("USER", "Gunakan filter tanggal / kategori untuk mempersempit tampilan.", None),
            ]),
            ("5. Melihat Rekam Jejak Tim (Atasan/HC)", [
                ("ATASAN/HC", "Klik menu CMP > Records lalu pilih tab Team View.", "B05-records-team.png"),
                ("ATASAN/HC", "Sistem menampilkan daftar bawahan sesuai hierarki organisasi.", None),
                ("ATASAN/HC", "Klik salah satu pekerja untuk melihat detail rekam jejak assessment per individu.", None),
            ]),
            ("6. Melihat Daftar & Download Sertifikat Kompetensi", [
                ("USER", "Klik menu CDP > Certification Management.", "B06-certification-management.png"),
                ("USER", "Sistem menampilkan daftar sertifikat pribadi beserta status (Aktif / Akan Expired / Expired).", None),
                ("USER", "Klik ikon unduh untuk mendapatkan file PDF sertifikat.", None),
                ("USER", "Sertifikat yang mendekati expired ditandai badge kuning; yang sudah expired badge merah.", None),
            ]),
            ("7. Analytics Dashboard (HC)", [
                ("HC", "Klik menu CMP > Analytics Dashboard.", "B07-analytics-dashboard.png"),
                ("HC", "Sistem menampilkan visualisasi: fail rate assessment, trend, skor elemen teknis, sertifikat expired, item analysis, dan gain score report.", None),
                ("HC", "Gunakan filter periode & bagian untuk analisis spesifik.", None),
                ("HC", "Ekspor data ke Excel untuk laporan manajemen.", None),
            ]),
        ]
    },
    "C. MODUL CDP — CAREER DEVELOPMENT": {
        "note": "Data IDP (track kompetensi, sub-kompetensi, deliverable) disusun dan di-upload oleh HC melalui menu HC > PROTON Data (lihat D.8). Role Coach, Coachee, dan Atasan hanya dapat melihat IDP sesuai lingkup akses masing-masing.",
        "items": [
            ("1. Melihat Individual Development Plan (IDP)", [
                ("COACHEE", "Klik menu CDP > Plan IDP.", "C01-plan-idp.png"),
                ("COACHEE", "Pilih Bagian, Unit, dan Track, lalu klik Muat Data. Sistem menampilkan silabus IDP sesuai track yang di-assign oleh HC.", None),
                ("COACHEE", "Klik salah satu deliverable untuk melihat detail progress (lihat C.3 Halaman Deliverable).", None),
                ("COACH", "Pada peran Coach, halaman Plan IDP menampilkan IDP seluruh coachee yang dipetakan. Gunakan filter untuk mempersempit tampilan.", None),
                ("ATASAN", "Pada peran Atasan, halaman Plan IDP menampilkan IDP bawahan langsung sesuai struktur organisasi.", None),
            ]),
            ("2. Sesi Coaching PROTON", [
                ("COACH", "Klik menu CDP > Proton Coaching.", "C02-coaching-proton.png"),
                ("COACH", "Sistem menampilkan daftar deliverable per coachee beserta status: Belum Upload / Pending Approval / Approved / Rejected.", None),
                ("COACH", "Pilih deliverable tertentu, klik Lihat Detail untuk masuk ke halaman progress.", None),
                ("COACH", "Input sesi coaching mengikuti kerangka PROTON 5 fase: Purpose, Realita, Options, To-Do, Outcome & Next step.", None),
                ("COACH", "Upload bukti coaching (foto kegiatan, dokumen kerja, hasil OJT) bila ada.", None),
                ("COACH", "Klik \"Simpan Sesi\". Sistem mencatat sesi ke histori & memicu update progress deliverable.", None),
            ]),
            ("3. Progress Deliverable — Lihat, Upload Evidence, Approval", [
                ("SEMUA", "Dari halaman Plan IDP atau Coaching Proton, klik nama deliverable > sistem membuka halaman Deliverable.", "C03-deliverable-detail.png"),
                ("SEMUA", "Halaman menampilkan: detail coachee & kompetensi, approval chain (Sr. Supervisor > Section Head > HC Review), dan riwayat status.", None),
                ("COACH", "Klik \"Upload Evidence\" > pilih file (PDF/JPG/PNG, max 10 MB) > klik \"Simpan\".", None),
                ("COACH", "Alternatif: \"Submit Evidence with Coaching\" — upload evidence sekaligus input sesi coaching dalam satu form.", None),
                ("ATASAN", "Pada deliverable yang sudah berisi evidence, Sr. Supervisor / Section Head klik \"Approve\" bila sesuai, atau \"Reject\" dengan mengisi alasan.", None),
                ("HC", "Lakukan HC Review sebagai approval final setelah reviewer memberikan approval tingkat-bagian.", None),
                ("SEMUA", "Histori seluruh aksi (upload, approve, reject, review) tercatat di bagian Riwayat Status halaman Deliverable.", None),
            ]),
            ("4. Melihat Histori Coaching PROTON", [
                ("USER", "Klik menu CDP > Histori Proton.", "C04-histori-proton.png"),
                ("USER", "Sistem menampilkan daftar pekerja beserta progress PROTON (Tahun 1, Tahun 2, Tahun 3) dan status (Lulus / Dalam Proses / Belum Mulai).", None),
                ("USER", "Klik \"Lihat Riwayat\" pada salah satu pekerja untuk melihat detail histori.", "C05-histori-proton-detail.png"),
                ("USER", "Halaman detail berisi data pekerja (Nama, NIP, Unit, Section, Jalur) serta timeline progress per tahun.", None),
            ]),
            ("5. Melihat Daftar Sertifikasi Pribadi (Certification Management)", [
                ("USER", "Klik menu CDP > Certification Management.", "B06-certification-management.png"),
                ("USER", "Sistem menampilkan ringkasan: Total Sertifikat, Aktif, Akan Expired, Expired.", None),
                ("USER", "Gunakan filter (Bagian, Unit, Status, Kategori) untuk mempersempit tampilan.", None),
                ("USER", "Klik salah satu baris atau ikon unduh untuk melihat detail dan mengunduh PDF sertifikat.", None),
            ]),
        ]
    },
    "D. MODUL HC — PENGELOLAAN DATA & PENGEMBANGAN KOMPETENSI": {
        "note": "Seluruh fitur di modul ini secara teknis dapat diakses oleh role Admin dan HC. Namun secara operasional, pengelolaan dilakukan oleh fungsi Human Capital (HC); peran Admin hanya untuk back-up teknis dan tidak dibahas pada TKI ini. Menu HC diakses melalui navbar \"Kelola Data\".",
        "items": [
            ("1. Kelola Data Pekerja", None, "D00-admin-dashboard.png"),
            ("1.1 Tambah Pekerja Manual", [
                ("HC", "Klik menu Kelola Data > Worker Management.", "D01-manage-workers.png"),
                ("HC", "Klik \"Tambah Pekerja\".", "D02-create-worker.png"),
                ("HC", "Isi form: Nama Lengkap, Email, NIP/Nopeg, Tanggal Bergabung, Jabatan, Directorate, Bagian, Unit.", None),
                ("HC", "Pilih role pekerja: Admin / HC / Direktur / VP / Manager / Section Head / Sr Supervisor / Coach / Supervisor / Coachee.", None),
                ("HC", "Tentukan password & konfirmasi password (minimal 6 karakter).", None),
                ("HC", "Klik \"Simpan Pekerja\".", None),
            ]),
            ("1.2 Import Pekerja Massal via Excel", [
                ("HC", "Pada halaman Worker Management, klik \"Import Excel\".", "D03-import-workers.png"),
                ("HC", "Klik \"Download Template\" untuk mendapatkan template Excel (lihat Lampiran II).", None),
                ("HC", "Isi template sesuai kolom yang disediakan: Nama, Email, NIP, Jabatan, Bagian, Unit, Directorate, Role, Tgl Bergabung, Password. Jangan ubah header.", None),
                ("HC", "Upload file Excel (drag & drop atau klik area upload).", None),
                ("HC", "Klik \"Proses Import\". Sistem melakukan validasi dan bulk insert dengan ringkasan hasil (berhasil / error / dilewati).", None),
            ]),
            ("1.3 Edit & Detail Pekerja", [
                ("HC", "Pada daftar pekerja, klik ikon Edit untuk mengubah data pekerja.", None),
                ("HC", "Edit field yang diperlukan (termasuk perubahan Role yang akan me-refresh akses pengguna) lalu klik Simpan.", None),
                ("HC", "Klik nama pekerja untuk masuk ke halaman Worker Detail berisi ringkasan lengkap beserta riwayat assessment & training.", None),
            ]),
            ("2. Struktur Organisasi", [
                ("HC", "Klik menu Kelola Data > Organization Structure.", "D04-manage-organization.png"),
                ("HC", "Sistem menampilkan pohon organisasi (Bagian > Unit) beserta ringkasan jumlah Bagian, Unit, dan status Aktif.", None),
                ("HC", "Klik panah pada node Bagian untuk expand daftar Unit di dalamnya.", None),
                ("HC", "Klik \"Tambah Unit\" untuk menambah Unit baru ke Bagian, atau ikon edit/hapus pada node untuk modifikasi.", None),
            ]),
            ("3. Pemetaan Coach ↔ Coachee", [
                ("HC", "Klik menu Kelola Data > Coach-Coachee Mapping.", "D05-coach-coachee-mapping.png"),
                ("HC", "Klik \"Tambah Mapping\" untuk assign coach ke coachee baru. Alternatif: import via Excel.", None),
                ("HC", "Pilih Seksi dan gunakan kolom Cari untuk mempersempit tampilan coach/coachee tertentu.", None),
                ("HC", "Aksi per mapping: Edit, Nonaktifkan, Graduated (tandai coachee telah selesai program).", None),
                ("HC", "Untuk memonitor beban coach, klik menu Kelola Data > Coach Workload.", "D06-coach-workload.png"),
                ("HC", "Sistem menampilkan jumlah coachee per coach beserta warning bila melampaui threshold (overloaded).", None),
                ("HC", "Klik \"Set Threshold\" untuk menyesuaikan batas maksimum coachee per coach.", None),
            ]),
            ("4. Kelola Kategori & Paket Soal Assessment", None),
            ("4.1 Kelola Kategori", [
                ("HC", "Klik menu Kelola Data > Assessment Categories.", "D07-manage-categories.png"),
                ("HC", "Sistem menampilkan daftar kategori beserta nilai lulus default (%) dan penandatangan sertifikat.", None),
                ("HC", "Klik \"Tambah Kategori\" untuk membuat kategori baru, atau ikon Edit/Hapus pada baris yang sudah ada.", None),
            ]),
            ("4.2 Kelola Paket Assessment", [
                ("HC", "Klik menu Kelola Data > Manage Assessment & Training.", "D08-manage-assessment.png"),
                ("HC", "Sistem menampilkan daftar assessment beserta kategori, jadwal, durasi, status (Open / Upcoming / InProgress / Closed), token, dan jumlah peserta.", None),
                ("HC", "Klik \"Buat Assessment\" > isi: nama paket, kategori, durasi, passing score, tipe (Online / Manual / PreTest / PostTest), tanggal buka/tutup.", None),
                ("HC", "Klik \"Simpan\".", None),
            ]),
            ("4.3 Import Soal via Excel", [
                ("HC", "Masuk ke detail paket assessment > klik \"Import Soal Excel\".", None),
                ("HC", "Download template sesuai tipe: Essay / Multiple Answer / Mix (lihat Lampiran I).", None),
                ("HC", "Isi template, upload, lakukan preview, lalu konfirmasi.", None),
                ("HC", "Klik \"Preview Paket\" untuk melihat tampilan akhir sebelum di-assign ke peserta.", None),
            ]),
            ("4.4 Penugasan Peserta (Assign Assessment)", [
                ("HC", "Pada detail paket, klik \"Assign Peserta\".", None),
                ("HC", "Pilih peserta via filter (bagian, jabatan) atau multi-select manual.", None),
                ("HC", "Klik \"Assign\". Peserta menerima notifikasi di HC Prime.", None),
            ]),
            ("5. Monitoring Assessment Realtime", [
                ("HC", "Klik menu Kelola Data > Assessment Monitoring.", "D09-assessment-monitoring.png"),
                ("HC", "Sistem menampilkan ringkasan: Grup Ditampilkan, Total Peserta, Selesai, Lulus.", None),
                ("HC", "Tabel monitoring berisi progress per assessment (jumlah peserta, jumlah selesai, jumlah lulus, progress bar).", None),
                ("HC", "Klik salah satu assessment untuk melihat detail per peserta: soal ke-n, sisa waktu, status realtime (via SignalR).", None),
                ("HC", "Untuk input hasil assessment manual (misal hasil interview), pakai tab Input Records pada halaman Manage Assessment & Training.", None),
            ]),
            ("6. Kelola Data Training Korporat", [
                ("HC", "Klik menu Kelola Data > Manage Assessment & Training > tab Input Records.", None),
                ("HC", "Klik \"Add Training\" untuk menambahkan record training secara manual, atau \"Import Training\" untuk import massal via Excel.", None),
                ("HC", "Edit record training peserta melalui ikon Edit. Upload sertifikat pendukung bila ada.", None),
            ]),
            ("7. Upload Dokumen KKJ & CPDP", [
                ("HC", "Klik menu Kelola Data > Kebutuhan Kompetensi Jabatan.", "D10-kkj-matrix.png"),
                ("HC", "Pilih tab bagian (RFCC / DHT-HMU / NGP / GAST).", None),
                ("HC", "Klik \"Upload File\" > pilih file KKJ (PDF/Excel, max 10 MB).", None),
                ("HC", "Sistem menyimpan dengan format nama file: yyyyMMddHHmmssfff_{GUID}_{namafile}.", None),
                ("HC", "Klik \"Riwayat File\" untuk melihat history versi dokumen.", None),
                ("HC", "Untuk CPDP: menu Kelola Data > CPDP File Management dengan alur serupa.", "D11-cpdp-files.png"),
            ]),
            ("8. PROTON Data — Setup IDP (Track, Kompetensi, Deliverable)", [
                ("HC", "Klik menu Kelola Data > Silabus & Coaching Guidance.", "D12-proton-data.png"),
                ("HC", "Pilih tab Status / Silabus / Coaching Guidance sesuai kebutuhan.", None),
                ("HC", "Pada tab Silabus, klik \"Import Silabus Excel\" untuk upload silabus baru.", None),
                ("HC", "Download template > isi dengan struktur: Track > Kompetensi > Sub-Kompetensi > Deliverable.", None),
                ("HC", "Upload file > preview > konfirmasi.", None),
                ("HC", "Untuk revisi entri tunggal, gunakan menu Kelola Data > Deliverable Progress Override.", None),
                ("HC", "Setelah data ter-upload, lakukan Track Assignment melalui mekanisme Coach-Coachee Mapping agar coachee tertentu terhubung ke track PROTON.", None),
            ]),
            ("9. Budget Training (Perencanaan Anggaran)", [
                ("HC", "Klik menu CMP > Budget Training (akses terbatas role HC/Admin).", "D13-budget-training.png"),
                ("HC", "Sistem menampilkan daftar anggaran training tahun berjalan beserta realisasi dan persentase serapan.", None),
                ("HC", "Klik \"Tambah\" untuk input budget baru: nama training, vendor, peserta target, estimasi biaya, tahun anggaran.", None),
                ("HC", "Alternatif: klik \"Import\" untuk input massal via Excel.", None),
                ("HC", "Update realisasi biaya via tombol inline, atau gunakan modal Bulk Update Realisasi.", None),
                ("HC", "Klik \"Export\" untuk mengunduh laporan anggaran dalam format Excel.", None),
            ]),
            ("10. Renewal Sertifikat", [
                ("HC", "Klik menu Kelola Data > Certificate Renewal.", "D14-renewal-certificate.png"),
                ("HC", "Sistem menampilkan ringkasan jumlah sertifikat Expired dan Akan Expired beserta daftar per assessment.", None),
                ("HC", "Gunakan filter Bagian, Unit, Kategori, Sub Kategori, Tipe, Status untuk mempersempit tampilan.", None),
                ("HC", "Klik grup sertifikat untuk expand daftar pekerja yang memiliki sertifikat tersebut.", None),
                ("HC", "Tandai sertifikat sebagai Renewal Processed setelah pekerja mengikuti training renewal.", None),
            ]),
        ]
    },
    "E. SISTEM & MONITORING": {
        "label": "TIDAK DISARANKAN DIMASUKKAN KE TKI FINAL",
        "note": "Bagian ini dipertahankan dalam outline sebagai referensi lengkap fungsionalitas website, namun seluruh isinya bersifat operasional teknis (audit log, maintenance mode, impersonation) yang merupakan ranah tim IT/Admin sistem, bukan prosedur kerja pengelolaan kompetensi oleh HC. Rekomendasi: saat kompilasi TKI Final, bagian E ini sebaiknya dipindahkan ke dokumen SOP IT terpisah atau dihapus dari TKI.",
        "items": [
            ("1. Notifikasi", [
                ("USER", "Pada navbar, ikon lonceng menampilkan jumlah notifikasi unread.", "E01-notification-dropdown.png"),
                ("USER", "Klik ikon untuk menampilkan dropdown daftar notifikasi terbaru (Assessment Selesai, Assessment Baru, Sertifikat Expired, dll.).", None),
                ("USER", "Klik salah satu item > sistem mengarah ke halaman terkait & menandai notifikasi sebagai read.", None),
                ("USER", "Klik \"Tandai semua dibaca\" untuk menandai seluruh notifikasi read sekaligus.", None),
            ]),
            ("2. Audit Log", [
                ("HC", "Klik menu Kelola Data > Audit Log.", "E02-audit-log.png"),
                ("HC", "Sistem menampilkan riwayat seluruh aksi admin/HC (assessment, worker, file, silabus, override, impersonation, dll.) beserta timestamp, user, aksi, dan deskripsi.", None),
                ("HC", "Gunakan filter Dari Tanggal & Sampai Tanggal, lalu klik \"Filter\" untuk mempersempit tampilan.", None),
                ("HC", "Klik \"Export Excel\" untuk mengunduh log dalam format Excel.", None),
            ]),
            ("3. Maintenance Mode", [
                ("HC", "Klik menu Kelola Data > Maintenance Mode.", "E03-maintenance.png"),
                ("HC", "Aktifkan toggle Mode Pemeliharaan untuk memblokir akses pengguna ke halaman yang dipilih.", None),
                ("HC", "Pilih cakupan: Seluruh Website atau Halaman Tertentu (pilih halaman per area: Home, CMP, CDP, Akun, Admin).", None),
                ("HC", "Isi pesan maintenance dan estimasi selesai. Admin dan HC tetap dapat mengakses halaman yang di-maintenance.", None),
                ("HC", "Klik \"Simpan Pengaturan\". Setelah selesai, toggle kembali ke nonaktif untuk mengembalikan akses normal.", None),
            ]),
            ("4. Impersonation (Troubleshooting)", [
                ("HC", "Klik menu Kelola Data, lalu pilih Impersonate.", "E04-impersonate.png"),
                ("HC", "Pilih View As HC / View As User untuk melihat portal sebagai peran tertentu, atau gunakan kolom pencarian untuk impersonate user spesifik (NIP/nama).", None),
                ("HC", "Mode impersonation bersifat read-only — tidak dapat mengubah data. Sesi maksimal 30 menit.", None),
                ("HC", "Banner impersonation tampil di top halaman. Klik Stop Impersonation untuk kembali ke sesi HC. Wajib stop impersonation setelah selesai troubleshoot.", None),
            ]),
        ]
    },
}

# =========================================================
# INDIKATOR KEBERHASILAN (BAB XI)
# =========================================================

INDIKATOR = [
    ("A. Adopsi Platform", [
        ">= 95% pekerja CSU Process (NGP, GAST, RFCC, DHT & HMU, Utilities II) login aktif minimal 1x per bulan.",
        "100% pekerja ISBL & OSBL memiliki akun HC Prime terdaftar dan teraktivasi.",
    ]),
    ("B. Assessment", [
        ">= 90% pekerja menyelesaikan assessment KKJ wajib tahunan sesuai jadwal.",
        "Rata-rata skor assessment >= 75% (passing score).",
        "<= 5% peserta mengalami disconnect/drop selama sesi assessment.",
        "100% soal Essay ter-grading manual oleh HC <= 14 hari setelah submit.",
    ]),
    ("C. IDP & Coaching PROTON", [
        "100% pekerja Production/Operations Team memiliki track IDP aktif di platform.",
        ">= 80% deliverable IDP per coachee tercatat progress-nya (upload evidence oleh Coach).",
        ">= 80% sesi coaching PROTON terdokumentasi di platform.",
        "Rata-rata >= 1 sesi coaching per coachee per bulan.",
        "Gap closure kompetensi teknis >= 20% per semester (perbandingan skor assessment sebelum-sesudah).",
        "Beban coach <= 10 coachee per coach.",
    ]),
    ("D. Sertifikasi & Training", [
        "0 sertifikat expired tanpa rencana renewal terdaftar di menu Renewal Certificate.",
        ">= 80% training mandatory (Safety for Refinery, Gas Tester, SUPREME, PSAIMS, ERP, Confined Space) ter-complete per pekerja sesuai periode.",
        ">= 90% sertifikat training ter-upload ke platform.",
    ]),
    ("E. Kualitas Data & Governance", [
        "Matriks KKJ terupdate setiap perubahan organisasi dalam <= 30 hari.",
        "Seluruh dokumen KKJ/CPDP per bagian memiliki history versi yang ter-track di platform.",
        "<= 1% kasus data mismatch antara HC Prime dengan sistem HC Core.",
    ]),
]

LAMPIRAN = [
    "Lampiran I — Template Excel Import Soal (Essay / Multiple Answer / Mix)",
    "Lampiran II — Template Excel Import Worker",
    "Lampiran III — Template Excel Import Training",
    "Lampiran IV — Template Excel Import Silabus PROTON",
    "Lampiran V — Template Excel Import Budget Training",
    "Lampiran VI — Flowchart Alur Assessment Online",
    "Lampiran VII — Flowchart Alur Coaching PROTON (5 fase) & Deliverable Approval",
    "Lampiran VIII — Matriks Akses per Role (HC / Atasan / Coach / Coachee / User)",
    "Lampiran IX — Daftar URL & Environment",
]

# =========================================================
# BUILD .DOCX (template TKI-GAST-003)
# =========================================================

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            bord = OxmlElement(f'w:{edge}')
            bord.set(qn('w:val'), 'single')
            bord.set(qn('w:sz'), str(kwargs[edge]))
            bord.set(qn('w:color'), '000000')
            tcBorders.append(bord)
    tcPr.append(tcBorders)

def add_header_table(section):
    """Header table ala TKI-GAST-003."""
    header = section.header
    # clear default paragraph
    for p in header.paragraphs:
        p.clear()
    tbl = header.add_table(rows=5, cols=2, width=Cm(17))
    tbl.autofit = False
    col_widths = (Cm(7), Cm(10))
    rows_data = [
        ("FUNGSI: HUMAN CAPITAL", "NO. DOKUMEN : C-     /KPB3200/2026-S9"),
        ("", "NOMOR REVISI : 0"),
        ("", "UNIQUE ID : TKI-HC-XXX-REV 0"),
        ("", "BERLAKU T.M.T :      /  /2026"),
        ("JUDUL: PROSEDUR PENGGUNAAN HC PRIME UNTUK PENGELOLAAN & PENGEMBANGAN KOMPETENSI PEKERJA", "HALAMAN          :  DARI"),
    ]
    for ridx, (left, right) in enumerate(rows_data):
        row = tbl.rows[ridx]
        for cidx, txt in enumerate((left, right)):
            cell = row.cells[cidx]
            cell.width = col_widths[cidx]
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            run.font.name = 'Arial'
            run.font.size = Pt(8)
            run.bold = True
            set_cell_border(cell, top=4, left=4, bottom=4, right=4)
    # subtitle line
    p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TATA KERJA INDIVIDU")
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(9)


def add_footer(section):
    footer = section.footer
    for p in footer.paragraphs:
        p.clear()
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Pertamina Confidential  ©  PT. Kilang Pertamina Balikpapan 2026")
    run.italic = True
    run.font.name = 'Arial'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def style_run(run, size=11, bold=False, italic=False, color=None):
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading_roman(doc, level, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    size_map = {1: 14, 2: 12, 3: 11}
    style_run(run, size=size_map.get(level, 11), bold=True)


def add_note(doc, note_text, label=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(6)
    if label:
        run = p.add_run(f"[ {label} ]  ")
        style_run(run, size=10, bold=True, color=RGBColor(0xB0, 0x20, 0x20))
    run = p.add_run(note_text)
    style_run(run, size=10, italic=True, color=RGBColor(0x55, 0x55, 0x55))


def add_step(doc, letter, role, step_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"{letter}. ")
    style_run(run, size=11)
    run = p.add_run(f"[{role}] ")
    style_run(run, size=10, bold=True, color=RGBColor(0x0D, 0x6E, 0xFD))
    run = p.add_run(step_text)
    style_run(run, size=11)


def add_screenshot(doc, filename):
    path = SHOT / filename
    if not path.exists():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5)
        run = p.add_run(f"[SCREENSHOT: {filename} — file tidak ditemukan]")
        style_run(run, size=9, italic=True, color=RGBColor(0xB0, 0x00, 0x00))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(15))


def letter(idx):
    """idx 0 -> a, 1 -> b ..."""
    return chr(ord('a') + idx)


def build_docx(out_path):
    doc = Document()

    # page setup
    for section in doc.sections:
        section.top_margin = Cm(3.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
        add_header_table(section)
        add_footer(section)

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # BAB X
    add_heading_roman(doc, 1, "X. INSTRUKSI KERJA")

    for section_title, body in OUTLINE.items():
        add_heading_roman(doc, 2, section_title)
        if body.get("label"):
            add_note(doc, body.get("note", ""), label=body["label"])
        elif body.get("note"):
            add_note(doc, body["note"])
        for entry in body["items"]:
            title = entry[0]
            steps = entry[1] if len(entry) > 1 else None
            extra_shot = entry[2] if len(entry) > 2 else None
            add_heading_roman(doc, 3, title)
            if extra_shot:
                add_screenshot(doc, extra_shot)
            if steps:
                for i, (role, txt, shot) in enumerate(steps):
                    add_step(doc, letter(i), role, txt)
                    if shot:
                        add_screenshot(doc, shot)

    # Page break before BAB XI
    doc.add_page_break()
    add_heading_roman(doc, 1, "XI. INDIKATOR & UKURAN KEBERHASILAN")
    for group_title, metrics in INDIKATOR:
        add_heading_roman(doc, 2, group_title)
        for i, m in enumerate(metrics, 1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.first_line_indent = Cm(-0.6)
            run = p.add_run(f"{i}. {m}")
            style_run(run, size=11)

    # Lampiran
    doc.add_page_break()
    add_heading_roman(doc, 1, "XIV. LAMPIRAN (Usulan)")
    for lmp in LAMPIRAN:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(f"• {lmp}")
        style_run(run, size=11)

    doc.save(out_path)
    print(f"[OK] .docx written: {out_path}")


# =========================================================
# BUILD .HTML preview
# =========================================================

HTML_STYLE = """
<style>
 body { font-family: 'Arial', sans-serif; max-width: 1100px; margin: 24px auto; padding: 0 24px; color: #222; line-height: 1.55; }
 h1 { border-bottom: 3px solid #0d6efd; padding-bottom: 6px; color: #0d6efd; margin-top: 40px; }
 h2 { border-left: 5px solid #0d6efd; padding-left: 12px; margin-top: 32px; color: #0a3a72; }
 h3 { color: #333; margin-top: 22px; }
 .note { background: #fff8e1; border-left: 4px solid #f9a825; padding: 10px 14px; margin: 12px 0; font-size: 0.93em; }
 .label-warn { background: #fee; border-left: 4px solid #c62828; padding: 10px 14px; margin: 12px 0; font-size: 0.93em; }
 .label-warn .chip { background: #c62828; color: white; padding: 2px 8px; border-radius: 4px; font-weight: bold; font-size: 0.85em; margin-right: 8px; }
 .step { margin: 4px 0 4px 20px; }
 .role { display: inline-block; background: #0d6efd; color: white; padding: 1px 8px; border-radius: 3px; font-size: 0.82em; font-weight: bold; margin-right: 6px; }
 .role.HC { background: #f9a825; color: #000; }
 .role.COACH { background: #2e7d32; }
 .role.COACHEE { background: #6a1b9a; }
 .role.ATASAN { background: #00838f; }
 .role.USER { background: #455a64; }
 .role.SEMUA { background: #222; }
 img { max-width: 960px; width: 100%; border: 1px solid #ddd; border-radius: 6px; margin: 10px 0 20px 20px; box-shadow: 0 2px 8px rgba(0,0,0,0.08); }
 .metric { margin: 4px 0 4px 20px; }
 hr.pagebreak { border: none; border-top: 2px dashed #ccc; margin: 40px 0; }
 .doc-header { background: #f4f6f8; border: 1px solid #ddd; padding: 14px; border-radius: 6px; font-size: 0.92em; margin-bottom: 24px; }
 .doc-header table { width: 100%; border-collapse: collapse; }
 .doc-header td { padding: 4px 8px; font-size: 0.9em; }
 .doc-header td:first-child { font-weight: bold; width: 25%; }

 /* Struktur role & akses */
 .struktur-section { background: #fafbfc; border: 1px solid #e0e4eb; border-radius: 8px; padding: 20px 24px; margin: 24px 0; }
 .struktur-section h2 { margin-top: 0; }
 table.role-table { width: 100%; border-collapse: collapse; margin: 12px 0 24px; font-size: 0.92em; }
 table.role-table th, table.role-table td { border: 1px solid #d0d7de; padding: 7px 10px; text-align: left; vertical-align: top; }
 table.role-table th { background: #eef2f7; font-weight: 600; color: #1a365d; }
 table.role-table tbody tr:nth-child(even) { background: #f8fafc; }
 .lvl-pill { display: inline-block; padding: 2px 8px; border-radius: 10px; font-size: 0.78em; font-weight: 700; color: white; min-width: 24px; text-align: center; }
 .lvl-1 { background: #c62828; }
 .lvl-2 { background: #ef6c00; }
 .lvl-3 { background: #2e7d32; }
 .lvl-4 { background: #00838f; }
 .lvl-5 { background: #1565c0; }
 .lvl-6 { background: #6a1b9a; }
 .view-pill { display: inline-block; padding: 1px 8px; border-radius: 3px; font-size: 0.78em; font-weight: 700; color: white; }
 .view-Admin { background: #c62828; }
 .view-HC { background: #f9a825; color: #000; }
 .view-Atasan { background: #00838f; }
 .view-Coach { background: #2e7d32; }
 .view-Coachee { background: #6a1b9a; }
 .access-yes { color: #1b5e20; font-weight: 700; text-align: center; }
 .access-read { color: #1565c0; font-weight: 600; text-align: center; }
 .access-no { color: #bbb; text-align: center; }
</style>
"""

def role_class(role):
    # handle combos like "ATASAN/HC"
    first = role.split('/')[0].split(' ')[0]
    return first


# =========================================================
# STRUKTUR PERAN & AKSES — body HTML (digunakan di build_html dan build_struktur_html)
# =========================================================
STRUKTUR_BODY_HTML = """
<div class="struktur-section">
<h2 style="margin-top:0">Struktur Peran & Akses Website HC Prime</h2>
<p style="font-size:0.93em;color:#555;margin-top:0">
  Referensi kontekstual sebelum masuk ke prosedur BAB X. Website HC Prime menggunakan <b>RBAC</b>
  (Role-Based Access Control) dengan 10 role yang dikelompokkan ke dalam 6 level hierarki.
  Setiap role memiliki lingkup akses dan <em>default view</em> berbeda.
</p>

<h3>1. Hierarki Role (Level 1–6)</h3>
<table class="role-table">
  <thead>
    <tr><th style="width:60px">Level</th><th style="width:150px">Role</th><th>Kelompok</th><th>Lingkup Akses</th></tr>
  </thead>
  <tbody>
    <tr><td><span class="lvl-pill lvl-1">1</span></td><td><b>Admin</b></td><td>System Administrator</td><td>Full access — seluruh fitur sistem, maintenance, audit, impersonation. <i>Back-up teknis.</i></td></tr>
    <tr><td><span class="lvl-pill lvl-2">2</span></td><td><b>HC</b></td><td>Human Capital</td><td>Full access — seluruh fitur CMP, CDP, Kelola Data. <i>Pengelola operasional utama.</i></td></tr>
    <tr><td><span class="lvl-pill lvl-3">3</span></td><td>Direktur</td><td rowspan="3">Management</td><td rowspan="3">Full access (read-weighted) — dashboard, seluruh tim; tidak mengelola data master.</td></tr>
    <tr><td><span class="lvl-pill lvl-3">3</span></td><td>VP</td></tr>
    <tr><td><span class="lvl-pill lvl-3">3</span></td><td>Manager</td></tr>
    <tr><td><span class="lvl-pill lvl-4">4</span></td><td>Section Head</td><td rowspan="2">Section Supervisory</td><td rowspan="2">Akses tingkat seksi — IDP & deliverable seluruh coachee dalam satu seksi, approval tingkat bagian.</td></tr>
    <tr><td><span class="lvl-pill lvl-4">4</span></td><td>Sr Supervisor</td></tr>
    <tr><td><span class="lvl-pill lvl-5">5</span></td><td><b>Coach</b></td><td rowspan="2">Coaching</td><td>Akses ke coachee yang di-map. Upload evidence, input sesi PROTON, lihat IDP coachee.</td></tr>
    <tr><td><span class="lvl-pill lvl-5">5</span></td><td>Supervisor</td><td>Akses view sama dengan Coach, tanpa coachee mapping.</td></tr>
    <tr><td><span class="lvl-pill lvl-6">6</span></td><td><b>Coachee</b></td><td>Operational</td><td>Akses ke data sendiri — IDP, assessment, records, sertifikat, histori coaching personal.</td></tr>
  </tbody>
</table>

<h3>2. Default View per Role</h3>
<p style="font-size:0.92em;color:#555;margin-top:0">
  Website memilih tampilan default berbeda untuk setiap role saat login. Ini menentukan widget dan menu yang paling menonjol di halaman Beranda.
</p>
<table class="role-table">
  <thead>
    <tr><th style="width:200px">Role</th><th style="width:140px">Default View</th><th>Karakteristik Tampilan</th></tr>
  </thead>
  <tbody>
    <tr><td>Admin</td><td><span class="view-pill view-Admin">Admin</span></td><td>Dashboard sistem, alert sertifikat expired, akses Kelola Data.</td></tr>
    <tr><td>HC</td><td><span class="view-pill view-HC">HC</span></td><td>Analytics dashboard, Kelola Data, ringkasan status pekerja lintas bagian.</td></tr>
    <tr><td>Direktur / VP / Manager / Section Head / Sr Supervisor</td><td><span class="view-pill view-Atasan">Atasan</span></td><td>Team view, progress bawahan, ringkasan assessment/coaching seksi.</td></tr>
    <tr><td>Coach / Supervisor</td><td><span class="view-pill view-Coach">Coach</span></td><td>Daftar coachee yang di-map, pending approval deliverable, histori coaching.</td></tr>
    <tr><td>Coachee (+ semua role lain)</td><td><span class="view-pill view-Coachee">Coachee</span></td><td>IDP pribadi, assessment ditugaskan, progress personal, sertifikat pribadi.</td></tr>
  </tbody>
</table>

<h3>3. Pemetaan Position (Jabatan) ↔ Role Default</h3>
<p style="font-size:0.92em;color:#555;margin-top:0">
  Position adalah nama jabatan struktural di PT KPB. Saat HC menambahkan pekerja baru, field <b>Role</b> dipilih terpisah dari <b>Position</b>. Tabel berikut adalah rekomendasi mapping default yang umum dipakai.
</p>
<table class="role-table">
  <thead>
    <tr><th>Position / Jabatan</th><th style="width:160px">Role Default</th><th>Contoh Fungsi</th></tr>
  </thead>
  <tbody>
    <tr><td>Direktur Operasi</td><td><span class="lvl-pill lvl-3">Direktur</span></td><td>Pimpinan Direktorat</td></tr>
    <tr><td>VP Refinery / VP HC & Corp. Services</td><td><span class="lvl-pill lvl-3">VP</span></td><td>Pimpinan Fungsi Level VP</td></tr>
    <tr><td>Manager Process / Manager HC</td><td><span class="lvl-pill lvl-3">Manager</span></td><td>Manager Fungsi</td></tr>
    <tr><td>Section Head (GAST, NGP, RFCC, DHT & HMU, Utilities II)</td><td><span class="lvl-pill lvl-4">Section Head</span></td><td>Kepala Seksi</td></tr>
    <tr><td>Sr Supervisor / Senior Supervisor Shift</td><td><span class="lvl-pill lvl-4">Sr Supervisor</span></td><td>Atasan langsung Shift Supervisor</td></tr>
    <tr><td>Shift Supervisor</td><td><span class="lvl-pill lvl-5">Coach</span> / <span class="lvl-pill lvl-5">Supervisor</span></td><td>Atasan langsung Operator/Panelman; bila di-map sebagai pembina coachee → Coach</td></tr>
    <tr><td>Panelman (Control Room Operator)</td><td><span class="lvl-pill lvl-6">Coachee</span></td><td>Operator DCS/HMI</td></tr>
    <tr><td>Operator (Field Operator)</td><td><span class="lvl-pill lvl-6">Coachee</span></td><td>Field operator per unit</td></tr>
    <tr><td>HC Staff (Jr./Sr. Analyst, Business Partner)</td><td><span class="lvl-pill lvl-2">HC</span></td><td>Pengelola platform kompetensi</td></tr>
    <tr><td>System Administrator (IT)</td><td><span class="lvl-pill lvl-1">Admin</span></td><td>Back-up teknis sistem</td></tr>
  </tbody>
</table>

<h3>4. Matriks Akses per Halaman</h3>
<p style="font-size:0.92em;color:#555;margin-top:0">
  ✅ = akses penuh (baca &amp; tulis) &nbsp;·&nbsp; 👁 = read-only (data sendiri/bawahan) &nbsp;·&nbsp; — = tidak punya akses
</p>
<table class="role-table">
  <thead>
    <tr>
      <th style="text-align:left">Halaman / Fitur</th>
      <th style="width:60px"><span class="view-pill view-Admin">Admin</span></th>
      <th style="width:60px"><span class="view-pill view-HC">HC</span></th>
      <th style="width:70px"><span class="view-pill view-Atasan">Atasan</span></th>
      <th style="width:70px"><span class="view-pill view-Coach">Coach</span></th>
      <th style="width:80px"><span class="view-pill view-Coachee">Coachee</span></th>
    </tr>
  </thead>
  <tbody>
    <tr><td colspan="6" style="background:#eef2f7;font-weight:700;color:#0d6efd">AKSES &amp; BERANDA</td></tr>
    <tr><td>Login / Beranda / Profil</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td></tr>

    <tr><td colspan="6" style="background:#eef2f7;font-weight:700;color:#0d6efd">MODUL CMP — Competency Management</td></tr>
    <tr><td>Dokumen KKJ — lihat</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td></tr>
    <tr><td>Ikut Assessment Online</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td></tr>
    <tr><td>Records (rekam jejak pribadi)</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td></tr>
    <tr><td>Records Team (rekam jejak tim)</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-read">👁</td><td class="access-read">👁</td><td class="access-no">—</td></tr>
    <tr><td>Certification Management</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-read">👁</td><td class="access-read">👁</td><td class="access-read">👁</td></tr>
    <tr><td>Analytics Dashboard</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>
    <tr><td>Budget Training</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>

    <tr><td colspan="6" style="background:#eef2f7;font-weight:700;color:#0d6efd">MODUL CDP — Career Development</td></tr>
    <tr><td>Plan IDP — lihat silabus</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-read">👁</td><td class="access-read">👁</td><td class="access-read">👁</td></tr>
    <tr><td>Coaching PROTON — input sesi &amp; upload evidence</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td></tr>
    <tr><td>Deliverable — approve / reject</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td></tr>
    <tr><td>Deliverable — HC Review final</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>
    <tr><td>Histori PROTON</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-read">👁</td><td class="access-read">👁</td><td class="access-read">👁</td></tr>

    <tr><td colspan="6" style="background:#eef2f7;font-weight:700;color:#0d6efd">MODUL HC — Pengelolaan Data</td></tr>
    <tr><td>Worker Management (CRUD + Import)</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>
    <tr><td>Organization Structure</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>
    <tr><td>Coach-Coachee Mapping &amp; Workload</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>
    <tr><td>Kategori &amp; Paket Assessment</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>
    <tr><td>Assessment Monitoring (real-time)</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>
    <tr><td>Upload Dokumen KKJ / CPDP</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>
    <tr><td>PROTON Data (silabus &amp; guidance)</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>
    <tr><td>Renewal Certificate</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>

    <tr><td colspan="6" style="background:#eef2f7;font-weight:700;color:#0d6efd">SISTEM &amp; MONITORING (non-TKI)</td></tr>
    <tr><td>Notifikasi</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td></tr>
    <tr><td>Audit Log</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>
    <tr><td>Maintenance Mode</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>
    <tr><td>Impersonation</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>
  </tbody>
</table>

<p style="font-size:0.85em;color:#666;font-style:italic;margin-top:8px">
  <b>Catatan:</b> Matriks ini hasil ekstraksi dari atribut <code>[Authorize(Roles=...)]</code> pada controller (<code>UserRoles.RolesCoachAndAbove</code>, <code>UserRoles.RolesReviewerAndAbove</code>) dan aturan <em>row-level scoping</em> (Coach hanya lihat coachee yang di-map; Atasan hanya lihat bawahan dalam seksi yang sama).
</p>
</div>
"""


def build_struktur_html(out_path):
    """Standalone file berisi hanya Struktur Peran & Akses."""
    html = [
        '<!doctype html><html lang="id"><head><meta charset="utf-8">',
        '<title>Struktur Peran & Akses — HC Prime</title>',
        HTML_STYLE,
        '</head><body>',
        """
    <div class="doc-header">
      <table>
        <tr><td>Dokumen</td><td>Struktur Peran &amp; Akses Website HC Prime</td></tr>
        <tr><td>Konteks</td><td>Referensi pendukung Draft TKI — PT Kilang Pertamina Balikpapan</td></tr>
        <tr><td>Status</td><td>Standalone preview</td></tr>
      </table>
    </div>
        """,
        STRUKTUR_BODY_HTML,
        "</body></html>",
    ]
    out_path.write_text("\n".join(html), encoding="utf-8")
    print(f"[OK] .html written: {out_path}")


def build_html(out_path):
    parts = ['<!doctype html><html lang="id"><head><meta charset="utf-8"><title>Draft BAB X INSTRUKSI KERJA - HC Prime TKI</title>', HTML_STYLE, '</head><body>']
    parts.append("""
    <div class="doc-header">
      <table>
        <tr><td>Fungsi</td><td>Human Capital</td></tr>
        <tr><td>No. Dokumen</td><td>C-     /KPB3200/2026-S9</td></tr>
        <tr><td>Unique ID</td><td>TKI-HC-XXX Rev 0</td></tr>
        <tr><td>Judul</td><td>PROSEDUR PENGGUNAAN HC PRIME UNTUK PENGELOLAAN & PENGEMBANGAN KOMPETENSI PEKERJA</td></tr>
        <tr><td>Status</td><td><b>Draft BAB X + BAB XI</b> — preview untuk review</td></tr>
      </table>
    </div>
    """)

    # ============================================================
    # STRUKTUR PERAN & AKSES (konten tambahan — preview HTML only)
    # ============================================================
    parts.append(STRUKTUR_BODY_HTML)
    # ===BEGIN_REMOVED_BLOCK===
    __removed = """
    <p style="font-size:0.93em;color:#555;margin-top:0">
      Referensi kontekstual sebelum masuk ke prosedur BAB X. Website HC Prime menggunakan <b>RBAC</b>
      (Role-Based Access Control) dengan 10 role yang dikelompokkan ke dalam 6 level hierarki.
      Setiap role memiliki lingkup akses dan <em>default view</em> berbeda.
    </p>

    <h3>1. Hierarki Role (Level 1–6)</h3>
    <table class="role-table">
      <thead>
        <tr><th style="width:60px">Level</th><th style="width:150px">Role</th><th>Kelompok</th><th>Lingkup Akses</th></tr>
      </thead>
      <tbody>
        <tr><td><span class="lvl-pill lvl-1">1</span></td><td><b>Admin</b></td><td>System Administrator</td><td>Full access — seluruh fitur sistem, maintenance, audit, impersonation. <i>Back-up teknis.</i></td></tr>
        <tr><td><span class="lvl-pill lvl-2">2</span></td><td><b>HC</b></td><td>Human Capital</td><td>Full access — seluruh fitur CMP, CDP, Kelola Data. <i>Pengelola operasional utama.</i></td></tr>
        <tr><td><span class="lvl-pill lvl-3">3</span></td><td>Direktur</td><td rowspan="3">Management</td><td rowspan="3">Full access (read-weighted) — dashboard, seluruh tim; tidak mengelola data master.</td></tr>
        <tr><td><span class="lvl-pill lvl-3">3</span></td><td>VP</td></tr>
        <tr><td><span class="lvl-pill lvl-3">3</span></td><td>Manager</td></tr>
        <tr><td><span class="lvl-pill lvl-4">4</span></td><td>Section Head</td><td rowspan="2">Section Supervisory</td><td rowspan="2">Akses tingkat seksi — IDP & deliverable seluruh coachee dalam satu seksi, approval tingkat bagian.</td></tr>
        <tr><td><span class="lvl-pill lvl-4">4</span></td><td>Sr Supervisor</td></tr>
        <tr><td><span class="lvl-pill lvl-5">5</span></td><td><b>Coach</b></td><td rowspan="2">Coaching</td><td>Akses ke coachee yang di-map. Upload evidence, input sesi PROTON, lihat IDP coachee.</td></tr>
        <tr><td><span class="lvl-pill lvl-5">5</span></td><td>Supervisor</td><td>Akses view sama dengan Coach, tanpa coachee mapping.</td></tr>
        <tr><td><span class="lvl-pill lvl-6">6</span></td><td><b>Coachee</b></td><td>Operational</td><td>Akses ke data sendiri — IDP, assessment, records, sertifikat, histori coaching personal.</td></tr>
      </tbody>
    </table>

    <h3>2. Default View per Role</h3>
    <p style="font-size:0.92em;color:#555;margin-top:0">
      Website memilih tampilan default berbeda untuk setiap role saat login. Ini menentukan widget dan menu yang paling menonjol di halaman Beranda.
    </p>
    <table class="role-table">
      <thead>
        <tr><th style="width:200px">Role</th><th style="width:140px">Default View</th><th>Karakteristik Tampilan</th></tr>
      </thead>
      <tbody>
        <tr><td>Admin</td><td><span class="view-pill view-Admin">Admin</span></td><td>Dashboard sistem, alert sertifikat expired, akses Kelola Data.</td></tr>
        <tr><td>HC</td><td><span class="view-pill view-HC">HC</span></td><td>Analytics dashboard, Kelola Data, ringkasan status pekerja lintas bagian.</td></tr>
        <tr><td>Direktur / VP / Manager / Section Head / Sr Supervisor</td><td><span class="view-pill view-Atasan">Atasan</span></td><td>Team view, progress bawahan, ringkasan assessment/coaching seksi.</td></tr>
        <tr><td>Coach / Supervisor</td><td><span class="view-pill view-Coach">Coach</span></td><td>Daftar coachee yang di-map, pending approval deliverable, histori coaching.</td></tr>
        <tr><td>Coachee (+ semua role lain)</td><td><span class="view-pill view-Coachee">Coachee</span></td><td>IDP pribadi, assessment ditugaskan, progress personal, sertifikat pribadi.</td></tr>
      </tbody>
    </table>

    <h3>3. Pemetaan Position (Jabatan) ↔ Role Default</h3>
    <p style="font-size:0.92em;color:#555;margin-top:0">
      Position adalah nama jabatan struktural di PT KPB. Saat HC menambahkan pekerja baru, field <b>Role</b> dipilih terpisah dari <b>Position</b>. Tabel berikut adalah rekomendasi mapping default yang umum dipakai.
    </p>
    <table class="role-table">
      <thead>
        <tr><th>Position / Jabatan</th><th style="width:160px">Role Default</th><th>Contoh Fungsi</th></tr>
      </thead>
      <tbody>
        <tr><td>Direktur Operasi</td><td><span class="lvl-pill lvl-3">Direktur</span></td><td>Pimpinan Direktorat</td></tr>
        <tr><td>VP Refinery / VP HC & Corp. Services</td><td><span class="lvl-pill lvl-3">VP</span></td><td>Pimpinan Fungsi Level VP</td></tr>
        <tr><td>Manager Process / Manager HC</td><td><span class="lvl-pill lvl-3">Manager</span></td><td>Manager Fungsi</td></tr>
        <tr><td>Section Head (GAST, NGP, RFCC, DHT & HMU, Utilities II)</td><td><span class="lvl-pill lvl-4">Section Head</span></td><td>Kepala Seksi</td></tr>
        <tr><td>Sr Supervisor / Senior Supervisor Shift</td><td><span class="lvl-pill lvl-4">Sr Supervisor</span></td><td>Atasan langsung Shift Supervisor</td></tr>
        <tr><td>Shift Supervisor</td><td><span class="lvl-pill lvl-5">Coach</span> / <span class="lvl-pill lvl-5">Supervisor</span></td><td>Atasan langsung Operator/Panelman; bila di-map sebagai pembina coachee → Coach</td></tr>
        <tr><td>Panelman (Control Room Operator)</td><td><span class="lvl-pill lvl-6">Coachee</span></td><td>Operator DCS/HMI</td></tr>
        <tr><td>Operator (Field Operator)</td><td><span class="lvl-pill lvl-6">Coachee</span></td><td>Field operator per unit</td></tr>
        <tr><td>HC Staff (Jr./Sr. Analyst, Business Partner)</td><td><span class="lvl-pill lvl-2">HC</span></td><td>Pengelola platform kompetensi</td></tr>
        <tr><td>System Administrator (IT)</td><td><span class="lvl-pill lvl-1">Admin</span></td><td>Back-up teknis sistem</td></tr>
      </tbody>
    </table>

    <h3>4. Matriks Akses per Halaman</h3>
    <p style="font-size:0.92em;color:#555;margin-top:0">
      ✅ = akses penuh (baca & tulis) &nbsp;·&nbsp; 👁 = read-only (data sendiri/bawahan) &nbsp;·&nbsp; — = tidak punya akses
    </p>
    <table class="role-table">
      <thead>
        <tr>
          <th style="text-align:left">Halaman / Fitur</th>
          <th style="width:60px"><span class="view-pill view-Admin">Admin</span></th>
          <th style="width:60px"><span class="view-pill view-HC">HC</span></th>
          <th style="width:70px"><span class="view-pill view-Atasan">Atasan</span></th>
          <th style="width:70px"><span class="view-pill view-Coach">Coach</span></th>
          <th style="width:80px"><span class="view-pill view-Coachee">Coachee</span></th>
        </tr>
      </thead>
      <tbody>
        <tr><td colspan="6" style="background:#eef2f7;font-weight:700;color:#0d6efd">AKSES & BERANDA</td></tr>
        <tr><td>Login / Beranda / Profil</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td></tr>

        <tr><td colspan="6" style="background:#eef2f7;font-weight:700;color:#0d6efd">MODUL CMP — Competency Management</td></tr>
        <tr><td>Dokumen KKJ — lihat</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td></tr>
        <tr><td>Ikut Assessment Online</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td></tr>
        <tr><td>Records (rekam jejak pribadi)</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td></tr>
        <tr><td>Records Team (rekam jejak tim)</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-read">👁</td><td class="access-read">👁</td><td class="access-no">—</td></tr>
        <tr><td>Certification Management</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-read">👁</td><td class="access-read">👁</td><td class="access-read">👁</td></tr>
        <tr><td>Analytics Dashboard</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>
        <tr><td>Budget Training</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>

        <tr><td colspan="6" style="background:#eef2f7;font-weight:700;color:#0d6efd">MODUL CDP — Career Development</td></tr>
        <tr><td>Plan IDP — lihat silabus</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-read">👁</td><td class="access-read">👁</td><td class="access-read">👁</td></tr>
        <tr><td>Coaching PROTON — input sesi & upload evidence</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td></tr>
        <tr><td>Deliverable — approve / reject</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td></tr>
        <tr><td>Deliverable — HC Review final</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>
        <tr><td>Histori PROTON</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-read">👁</td><td class="access-read">👁</td><td class="access-read">👁</td></tr>

        <tr><td colspan="6" style="background:#eef2f7;font-weight:700;color:#0d6efd">MODUL HC — Pengelolaan Data</td></tr>
        <tr><td>Worker Management (CRUD + Import)</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>
        <tr><td>Organization Structure</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>
        <tr><td>Coach-Coachee Mapping & Workload</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>
        <tr><td>Kategori & Paket Assessment</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>
        <tr><td>Assessment Monitoring (real-time)</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>
        <tr><td>Upload Dokumen KKJ / CPDP</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>
        <tr><td>PROTON Data (silabus & guidance)</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>
        <tr><td>Renewal Certificate</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>

        <tr><td colspan="6" style="background:#eef2f7;font-weight:700;color:#0d6efd">SISTEM & MONITORING (non-TKI)</td></tr>
        <tr><td>Notifikasi</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-yes">✅</td></tr>
        <tr><td>Audit Log</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>
        <tr><td>Maintenance Mode</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>
        <tr><td>Impersonation</td><td class="access-yes">✅</td><td class="access-yes">✅</td><td class="access-no">—</td><td class="access-no">—</td><td class="access-no">—</td></tr>
      </tbody>
    </table>

    <p style="font-size:0.85em;color:#666;font-style:italic;margin-top:8px">
      <b>Catatan:</b> Matriks ini hasil ekstraksi dari atribut <code>[Authorize(Roles=...)]</code> pada controller (<code>UserRoles.RolesCoachAndAbove</code>, <code>UserRoles.RolesReviewerAndAbove</code>) dan aturan <em>row-level scoping</em> (Coach hanya lihat coachee yang di-map; Atasan hanya lihat bawahan dalam seksi yang sama).
    </p>
    </div>
    """
    # ===END_REMOVED_BLOCK===

    parts.append("<h1>X. INSTRUKSI KERJA</h1>")
    for section_title, body in OUTLINE.items():
        parts.append(f"<h2>{section_title}</h2>")
        if body.get("label"):
            parts.append(f'<div class="label-warn"><span class="chip">LABEL: {body["label"]}</span>{body.get("note","")}</div>')
        elif body.get("note"):
            parts.append(f'<div class="note">{body["note"]}</div>')
        for entry in body["items"]:
            title = entry[0]
            steps = entry[1] if len(entry) > 1 else None
            extra_shot = entry[2] if len(entry) > 2 else None
            parts.append(f"<h3>{title}</h3>")
            if extra_shot and (SHOT / extra_shot).exists():
                parts.append(f'<img src="screenshots/{extra_shot}" alt="{extra_shot}">')
            if steps:
                for i, (role, txt, shot) in enumerate(steps):
                    rc = role_class(role)
                    parts.append(f'<div class="step">{letter(i)}. <span class="role {rc}">{role}</span>{txt}</div>')
                    if shot and (SHOT / shot).exists():
                        parts.append(f'<img src="screenshots/{shot}" alt="{shot}">')

    parts.append('<hr class="pagebreak">')
    parts.append("<h1>XI. INDIKATOR & UKURAN KEBERHASILAN</h1>")
    for group_title, metrics in INDIKATOR:
        parts.append(f"<h2>{group_title}</h2>")
        for i, m in enumerate(metrics, 1):
            parts.append(f'<div class="metric">{i}. {m}</div>')

    parts.append('<hr class="pagebreak">')
    parts.append("<h1>XIV. LAMPIRAN (Usulan)</h1><ul>")
    for lmp in LAMPIRAN:
        parts.append(f"<li>{lmp}</li>")
    parts.append("</ul>")

    parts.append("</body></html>")
    out_path.write_text("\n".join(parts), encoding="utf-8")
    print(f"[OK] .html written: {out_path}")


if __name__ == "__main__":
    build_docx(HERE / "Draft-BAB-X-INSTRUKSI-KERJA.docx")
    build_html(HERE / "Draft-BAB-X-INSTRUKSI-KERJA.html")
    build_struktur_html(HERE / "Struktur-Peran-Akses-HC-Prime.html")
